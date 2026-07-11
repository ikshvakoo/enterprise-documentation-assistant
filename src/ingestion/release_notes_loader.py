from pathlib import Path

from docx import Document
from pypdf import PdfReader

from models import SourceDocument
from text_utils import clean_text
from text_utils import join_nonempty


def load_release_notes(directory: Path) -> list[SourceDocument]:
    """Load all release-note files from a directory into SourceDocument objects."""
    documents: list[SourceDocument] = []
    for path in sorted(directory.glob("*.pdf")):
        documents.extend(load_pdf(path, "release_note"))
    for path in sorted(directory.glob("*.docx")):
        documents.extend(load_docx(path, "release_note"))
    return documents


def load_pdf(path: Path, source_type: str) -> list[SourceDocument]:
    """Extract text from a PDF release note, creating one document per page."""
    reader = PdfReader(str(path))
    documents: list[SourceDocument] = []
    for page_number, page in enumerate(reader.pages, start=1):
        # Page-level documents keep citations specific enough for users to
        # verify an answer against the original release note.
        text = clean_text(page.extract_text() or "")
        if not text:
            continue
        documents.append(
            SourceDocument(
                doc_id=f"{source_type}:{path.name}:page:{page_number}",
                source_type=source_type,
                title=f"{path.stem} - page {page_number}",
                body=text,
                source_path=str(path),
                metadata={"file_name": path.name, "page": page_number},
            )
        )
    return documents


def load_docx(path: Path, source_type: str) -> list[SourceDocument]:
    """Extract paragraphs and table text from a DOCX release-note document."""
    document = Document(path)
    parts: list[str] = []
    for paragraph in document.paragraphs:
        if paragraph.text.strip():
            parts.append(paragraph.text)
    # Tables often contain release-note details; include row text with separators
    # so the search index can still match content inside cells.
    for table in document.tables:
        for row in table.rows:
            values = [cell.text for cell in row.cells if cell.text.strip()]
            if values:
                parts.append(" | ".join(values))

    body = join_nonempty(parts)
    if not body:
        return []

    return [
        SourceDocument(
            doc_id=f"{source_type}:{path.name}",
            source_type=source_type,
            title=path.stem,
            body=body,
            source_path=str(path),
            metadata={"file_name": path.name},
        )
    ]
