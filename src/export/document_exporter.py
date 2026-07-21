"""Export generated documentation drafts to common file formats."""

from pathlib import Path

from docx import Document


def write_markdown(content: str, path: Path) -> None:
    """Write Markdown content to disk using UTF-8 encoding."""
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(content, encoding="utf-8", newline="\n")


def write_release_notes_docx(markdown: str, path: Path) -> None:
    """Convert the generated Markdown-like release notes into a DOCX file.

    This converter supports the subset of Markdown produced by the local
    generator: headings, bullets, numbered evidence items, and traceability
    tables. Keeping the converter small makes it predictable for the demo.
    """
    path.parent.mkdir(parents=True, exist_ok=True)
    document = Document()
    table_buffer: list[list[str]] = []

    for raw_line in markdown.splitlines():
        line = raw_line.strip()
        if not line:
            flush_table(document, table_buffer)
            table_buffer.clear()
            continue

        if line.startswith("|"):
            cells = [cell.strip() for cell in line.strip("|").split("|")]
            if not all(set(cell) <= {"-", " "} for cell in cells):
                table_buffer.append(cells)
            continue

        flush_table(document, table_buffer)
        table_buffer.clear()

        if line.startswith("# "):
            document.add_heading(line[2:].strip(), level=1)
        elif line.startswith("## "):
            document.add_heading(line[3:].strip(), level=2)
        elif line.startswith("### "):
            document.add_heading(line[4:].strip(), level=3)
        elif line.startswith("- "):
            document.add_paragraph(line[2:].strip(), style="List Bullet")
        elif looks_numbered(line):
            _, text = line.split(".", 1)
            document.add_paragraph(text.strip(), style="List Number")
        else:
            document.add_paragraph(line)

    flush_table(document, table_buffer)
    document.save(path)


def flush_table(document: Document, rows: list[list[str]]) -> None:
    """Write buffered Markdown table rows into the DOCX document."""
    if not rows:
        return
    column_count = max(len(row) for row in rows)
    table = document.add_table(rows=len(rows), cols=column_count)
    table.style = "Table Grid"
    for row_index, row in enumerate(rows):
        for column_index in range(column_count):
            table.cell(row_index, column_index).text = row[column_index] if column_index < len(row) else ""


def looks_numbered(line: str) -> bool:
    """Return True when a line starts with a simple numbered-list prefix."""
    prefix, separator, _ = line.partition(".")
    return bool(separator) and prefix.isdigit()
